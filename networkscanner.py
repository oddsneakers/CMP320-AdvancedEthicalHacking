# Import Libraries
import argparse
from prettytable import PrettyTable
from scapy.all import ARP, Ether, srp
import nmap
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

# Function to perform an ARP scan on a target IP address
def arp_scan(target_ip):
    print("Performing ARP Scan...")
    ether = Ether(dst="ff:ff:ff:ff:ff:ff")  # Set Ethernet frame destination as broadcast
    arp = ARP(pdst=target_ip)  # Create ARP request for the target IP
    result = srp(ether/arp, timeout=3, verbose=False)[0]  # Send the packet and get response

    ip_addresses = [received.psrc for sent, received in result]  # Extract IP addresses from responses
    print("IP Addresses found:")
    for ip in ip_addresses:
        print(ip)
    return ip_addresses

# Function to perform an Nmap scan on discovered IP addresses and output results to a document
def nmap_scan_and_output(ip_addresses):
    nm = nmap.PortScanner()  # Create a new PortScanner object
    document = Document()  # Create a new Word document
    document.add_heading('Vulnerability Assessment Report', 0).alignment = 1  # Add main heading
    document.add_page_break()  # Add a page break
    document.add_heading('Table of Contents', level=1)  # Add a table of contents heading
    document.add_paragraph('Contents here...\nNote: Update this section manually or with an automated tool.')
    document.add_page_break()  # Add another page break

    for ip in ip_addresses:
        print(f"Scanning {ip} with Nmap Full Scan and Vulnerability Script...")
        nm.scan(ip, arguments='-p 1-65535 -sV --script=vuln')  # Perform an Nmap scan
        output_filename = f"nmap_full_vuln_scan_{ip}.txt"
        extract_service_info_to_file(ip, nm, output_filename)  # Extract info to text file
        extract_service_info_to_document(ip, nm, document)  # Extract info to Word document

    report_filename = "vulnerability_assessment_report.docx"
    document.save(report_filename)  # Save the Word document
    print(f"White paper report generated: {report_filename}")

# Function to extract service information to a text file
def extract_service_info_to_file(ip, nm, filename):
    table = PrettyTable()  # Create a new PrettyTable object
    table.field_names = ["IP", "PROTOCOL", "PORT", "STATE", "SERVICE", "VERSION", "VULNERABILITIES"]
    populate_table(ip, nm, table)  # Populate table with scan results
    with open(filename, 'w') as file:
        file.write(str(table))  # Write table to file

# Function to extract service information to a Word document
def extract_service_info_to_document(ip, nm, document):
    document.add_heading(f"Nmap Scan Results", level=1)  # Add a heading for Nmap results
    document.add_heading(f"{ip}", level=2)  # Add a subheading for the IP address
    table = document.add_table(rows=1, cols=3)  # Create a new table in the document
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "PROTOCOL OPEN PORT"
    hdr_cells[1].text = "SERVICE"
    hdr_cells[2].text = "VERSION"
    for host in nm.all_hosts():
        for proto in nm[host].all_protocols():
            lport = nm[host][proto].keys()
            for port in sorted(lport):
                port_info = nm[host][proto][port]
                state = port_info['state']
                if state == 'open':
                    service = port_info['name']
                    version = port_info.get('product', '') + " " + port_info.get('version', '').strip()
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"{proto.upper()} {port}"
                    row_cells[1].text = service
                    row_cells[2].text = version

# Function to populate the PrettyTable with data from Nmap scan results
def populate_table(ip, nm, table):
    for host in nm.all_hosts():
        for proto in nm[host].all_protocols():
            lport = nm[host][proto].keys()
            for port in sorted(lport):
                port_info = nm[host][proto][port]
                service = port_info['name']
                state = port_info['state']
                version = port_info.get('product', '') + " " + port_info.get('version', '').strip()
                script_info = port_info.get('script', '')
                vuln_info = script_info.get('vuln', '') if script_info else 'N/A'
                if state == 'open':
                    table.add_row([ip, proto.upper(), str(port), state, service, version, vuln_info])

# Main execution block
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Network Vulnerability Scanner Tool')
    parser.add_argument('target_network', help='IP range to scan, e.g., 192.168.1.0/24')
    args = parser.parse_args()
    
    discovered_ips = arp_scan(args.target_network)
    if discovered_ips:
        nmap_scan_and_output(discovered_ips)
    else:
        print("No IP addresses discovered.")
